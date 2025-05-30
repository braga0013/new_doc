import whisper
import cv2
import os
import google.generativeai as genai
from fastapi import APIRouter, Query, File, UploadFile, Form
from fastapi.responses import FileResponse
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import tempfile
import shutil


genai.configure(api_key="")

router = APIRouter()

def transcrever_audio(video_path):
    model = whisper.load_model("base")
    result = model.transcribe(video_path)
    return result["text"]

def extrair_frames(video_path, output_dir="frames_extraidos"):
    cap = cv2.VideoCapture(video_path)
    success, frame = cap.read()
    count = 0
    frames_salvos = []
    os.makedirs(output_dir, exist_ok=True)
    while success:
        frame_path = os.path.join(output_dir, f"frame{count}.jpg")
        cv2.imwrite(frame_path, frame)
        frames_salvos.append(frame_path)
        for _ in range(90):
            success, frame = cap.read()
            if not success:
                break
        count += 1
    cap.release()
    return frames_salvos

def capturar_html_da_pagina(url):
    options = Options()
    options.add_argument('--headless')
    options.add_argument('--disable-gpu')
    options.add_argument('--no-sandbox')
    options.binary_location = "/usr/bin/chromium"
    driver = webdriver.Chrome(options=options)
    driver.get(url)
    driver.implicitly_wait(10)
    html = driver.page_source
    driver.quit()
    return html

def perguntar_ao_gemini(texto_base, pergunta_usuario):
    model = genai.GenerativeModel('models/gemini-1.5-flash-001')
    resposta = model.generate_content([texto_base, pergunta_usuario])
    return resposta.text

def salvar_docx(texto, caminho_arquivo):
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    for linha in texto.split('\n'):
        linha = linha.strip()

        # Títulos
        if linha.startswith('#### '):
            p = doc.add_heading(linha.replace('#### ', ''), level=4)
        elif linha.startswith('### '):
            p = doc.add_heading(linha.replace('### ', ''), level=3)
        elif linha.startswith('## '):
            p = doc.add_heading(linha.replace('## ', ''), level=2)
        elif linha.startswith('# '):
            p = doc.add_heading(linha.replace('# ', ''), level=1)

        # Lista com negrito parcial
        elif linha.startswith('* ') or linha.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            partes = linha[2:].split('**')
            for i, parte in enumerate(partes):
                run = p.add_run(parte)
                if i % 2 == 1:
                    run.bold = True

        # Negrito total
        elif linha.startswith('**') and linha.endswith('**') and len(linha) > 4:
            p = doc.add_paragraph()
            run = p.add_run(linha[2:-2])
            run.bold = True

        # Negrito parcial no corpo do texto
        elif '**' in linha:
            p = doc.add_paragraph()
            partes = linha.split('**')
            for i, parte in enumerate(partes):
                run = p.add_run(parte)
                if i % 2 == 1:
                    run.bold = True

        # Normal
        else:
            doc.add_paragraph(linha)

    doc.save(caminho_arquivo)

@router.post("/documentacao-automatica")
async def documentacao_automatica(
    video: UploadFile = File(..., description="Envie o vídeo para análise"),
    url: str = Form(None, description="URL da página para capturar HTML (opcional)"),
    detalhes: str = Form("", description="Detalhes específicos para serem considerados na documentação (opcional)"),
    pergunta: str = Form(
        "Analise cuidadosamente o vídeo, o código HTML fornecido e os detalhes adicionais, e gere uma documentação técnica completa e detalhada da página apresentada. "
        "Descreva todos os elementos interativos, fluxos de navegação, possíveis workflows, regras de negócio, exemplos práticos de uso, automações, integrações e aspectos técnicos relevantes.",
        description="Prompt para gerar documentação completa da interface com base em vídeo, HTML e detalhes extras."
    ),
    download: bool = Form(False, description="Se True, retorna um arquivo .docx para download")
):
    # Salvar vídeo temporariamente
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as tmp:
        shutil.copyfileobj(video.file, tmp)
        video_path = tmp.name

    texto_fala = transcrever_audio(video_path)
    frames_salvos = extrair_frames(video_path)
    descricao_visual = "Frames extraídos do vídeo: " + ", ".join(frames_salvos)

    texto_base = f"""
Conteúdo falado no vídeo:
{texto_fala}

Conteúdo visual do vídeo:
{descricao_visual}
"""

    if url:
        try:
            html = capturar_html_da_pagina(url)
            texto_base += f"\nCódigo HTML da página:\n{html}\n"
        except Exception as e:
            os.unlink(video_path)
            return {"erro": f"Erro ao capturar HTML: {str(e)}"}

    if detalhes:
        texto_base += f"\nDetalhes específicos fornecidos pelo usuário:\n{detalhes}\n"

    resposta = perguntar_ao_gemini(texto_base, pergunta)

    if download:
        caminho_docx = "documentacao_gerada.docx"
        salvar_docx(resposta, caminho_docx)
        os.unlink(video_path)
        return FileResponse(
            caminho_docx,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="documentacao_gerada.docx"
        )

    os.unlink(video_path)
    return {"resposta": resposta}